from docx import Document
import re
import time

path = "初中历史基础知识点.docx"

document = Document(path)

# 根据颜色获取类型
def get_run_type(run):
    rgb = run.font.color.rgb
    if not rgb:
        return "ques"
    else:
        return "ans"

# 获取段落是否为一道题
def isQuestion(paragraph):
    # 以标号开头
    if paragraph.style.name == "List Paragraph":
        return True
    # 以数字加"."或数字加"、"开头的文本
    pattern = r'^\d+(\.|、)'
    return bool(re.match(pattern, paragraph.text))

def isTitle(paragraph):
    # 标题格式
    if paragraph.style.name == "Heading 1":
        return True
    return False

result = []
quesList = None
question = ""
answer = ""
for i in range(len(document.paragraphs)):
    paragraph = document.paragraphs[i]

    if isTitle(paragraph):
        if quesList:
            result.append(quesList)
        quesList = [paragraph.text]
        continue

    for j in range(len(paragraph.runs)):
        run = paragraph.runs[j]
        text = run.text

        now_type = get_run_type(run)

        if now_type == "ques":
            question += text
        else:
            answer += text

            next_run = paragraph.runs[j+1] if j+1 < len(paragraph.runs) else None
            next_type = get_run_type(next_run) if next_run else None

            if next_type == "ques":
                question += '____'
                answer += ';'
            
            if not next_run:
                question += '____'

    # print(paragraph.style.name,isQuestion(paragraph))
    next_paragraph = document.paragraphs[i+1] if i+1 < len(document.paragraphs) else None
    if next_paragraph and not isQuestion(next_paragraph):
        continue

    #替换连续空格
    question = re.sub(' +', ' ', question)
    answer = re.sub(' +', ' ', answer)

    # 去除 \t \n 等字符
    question = question.replace('\\t', '').replace('\\n', '').replace('\t', '').replace('\n', '')
    answer = answer.replace('\\t', '').replace('\\n', '').replace('\t', '').replace('\n', '')

    # 去除数字开头的标号
    pattern = r'^\d+(\.|、)'
    if bool(re.match(pattern, question)):
        clean_text = re.sub(pattern, '', question)  
        question = clean_text.strip()    

    print(question)
    print(answer)
    quesList.append([question, answer])

    # 置空
    question = ""
    answer = ""

    print("======")
    # if i == 100:break
    
result.append(quesList)

from pprint import pprint
pprint(result)

lastId = 2427
lastBookId = 6
# 生成sql文件
i = 1
nameList = []
id = 0
for quesList in result:
    name = quesList[0]
    nameList.append(name)

    sql_result = ""

    for ques, ans in quesList[1:]:
        ansList = ans.split('\\n')

        _id = id + 1 + lastId
        bookId = lastBookId + i
        cardFront = ques.replace('"', '”')
        cardBehind = ' '.join(ansList).replace('"', '”')
        createTime = int(time.time())

        cardBehindMk = '<html><body>'
        for ans in ansList:
            cardBehindMk += '<p>' + ans + '</p>'
        cardBehindMk += '</body></html>'
        cardBehindMk = cardBehindMk.replace('\\n', '<br>').replace('"', '”')
        
        sql = f"""INSERT INTO "main"."Card" ("_id", "bookId", "cardFront", "cardBehind", "cardBehindMk", "createTime", "cardType", "repetitions", "interval", "easinessFactor", "lapsesCount", "stepsIndex", "isStar", "isChop", "templateType", "popupCount", "learnSort", "lastNewTime", "lastRevTime") VALUES ({_id},{bookId},"{cardFront}","{cardBehind}","{cardBehindMk}",{createTime},1,0,0.0,2.5,0,0,0,0,1,0,0,946656000000,946656000000);"""
        
        sql_result += sql + '\n'

        id += 1
    i += 1

    with open(name+'.sql', 'w', encoding='utf-8') as f:
        f.write(sql_result)
    
i = 1
sql_result = ""
for name in nameList:
    createTime = int(time.time())
    
    sql = f"""INSERT INTO "Book" ("_id", "mid", "title", "introduce", "cardCount", "categorys", "createTime", "updateTime", "learnSort", "cardTestSort", "lastPopupType", "bLastPopupType", "mLastPopupType", "planConfig") VALUES ({lastBookId + i}, NULL, '{name}', '{name}', 0, NULL, {createTime}, {createTime}, -1, '复习优先', 'NEW', 'NEW', 'NEW', 'd�3b��');"""
    sql_result += sql + '\n'
    i += 1
    
with open('book.sql', 'w', encoding='utf-8') as f:
    f.write(sql_result)
